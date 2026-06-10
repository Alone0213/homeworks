# -*- coding: utf-8 -*-
"""
生成家具城进销存管理系统 SQL 实现文档 (.docx)
按模块分章节，每个部分配解释表格 + 完整SQL代码
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式设置 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.25
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 辅助函数 ──
def add_code_block(doc, code_text, font_size=9):
    """添加带灰底的代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # 灰底
    shd = run._element.get_or_add_rPr()
    shdXml = shd.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    shd.append(shdXml)
    return p

def add_explain_table(doc, rows_data, col_widths=None):
    """添加解释表格: 自动适配列数"""
    if not rows_data:
        return None
    ncols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 数据行
    for i, row_vals in enumerate(rows_data):
        cells = table.rows[i].cells
        for j, val in enumerate(row_vals):
            cells[j].text = str(val)
            for p in cells[j].paragraphs:
                for run in p.runs:
                    if i == 0:  # 首行加粗作为表头
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 间距
    return table

def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


# ================================================================
# 封面
# ================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(120)
run = title_p.add_run('家具城进销存管理系统')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('SQL 数据库实现文档\n（含完整代码与表格解析）')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.add_run('技术栈：MySQL 8.0+ | InnoDB | 外键约束 | 触发器 | 存储过程 | 视图\n'
               '表数量：7 张 | 外键：6 个 | 触发器：6 个 | 视图：3 个\n'
               '2025年7月').font.size = Pt(11)

doc.add_page_break()

# ================================================================
# 目录页（手动）
# ================================================================
doc.add_heading('目  录', level=1)
toc_items = [
    '一、数据库设计总览 —— 7张表的关联关系',
    '二、建库与建表语句（DDL）',
    '    2.1 家具类型表 (FurnitureCategory)',
    '    2.2 供应商信息表 (Supplier)',
    '    2.3 客户信息表 (Customer)',
    '    2.4 家具信息表 (Furniture)',
    '    2.5 入库记录表 (WarehouseEntry)',
    '    2.6 销售记录表 (Sales)',
    '    2.7 收款记录表 (Receipt)',
    '三、参照完整性约束（外键）',
    '四、触发器（自动更新库存）',
    '    4.1 入库触发器 (3个)',
    '    4.2 销售触发器 (3个)',
    '五、存储过程（时间段统计）',
    '六、视图封装（简化查询）',
    '七、测试数据（预置记录）',
    '八、Web 应用功能说明（Flask 前端特性）',
    '九、常用查询示例',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ================================================================
# 一、数据库设计总览
# ================================================================
doc.add_heading('一、数据库设计总览 —— 7张表的关联关系', level=1)

doc.add_paragraph('本系统包含 7 张数据表，通过外键构成完整的引用网络。下图展示各表之间的关联关系：')

# 表关联说明表格
add_explain_table(doc, [
    ('父表 (1端)', '子表 (N端)', '外键字段', '关系说明'),
    ('FurnitureCategory', 'Furniture', 'CategoryID', '每种类型对应多件家具'),
    ('Supplier', 'Furniture', 'SupplierID', '每个供应商供应多件家具'),
    ('Furniture', 'WarehouseEntry', 'FurnitureID', '每件家具可有多次入库'),
    ('Furniture', 'Sales', 'FurnitureID', '每件家具可多次销售'),
    ('Customer', 'Sales', 'CustomerID', '每位客户可有多次购买'),
    ('Sales', 'Receipt', 'SaleID', '每笔销售可多次收款'),
])

doc.add_paragraph('ER 关系图（文字版）:')

# 画一个简单的ASCII图
ascii_er = """  FurnitureCategory ──1:N── Furniture ──1:N── WarehouseEntry
       │                    │
       │                    └──1:N── Sales ──1:N── Receipt
       │                             │
  Supplier ────1:N─────────────┘     │
                                     │
                               Customer ──1:N────────┘"""
p = doc.add_paragraph()
run = p.add_run(ascii_er)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ================================================================
# 二、建库与建表语句
# ================================================================
doc.add_heading('二、建库与建表语句（DDL）', level=1)

doc.add_paragraph('首先创建数据库，然后依次创建 7 张数据表。每张表均使用 InnoDB 引擎以支持外键约束和事务。')

# 建库
doc.add_heading('2.0 创建数据库', level=2)
add_explain_table(doc, [
    ('语句', 'CREATE DATABASE FurnitureDB'),
    ('字符集', 'utf8mb4 —— 支持中文、表情符号'),
    ('排序规则', 'utf8mb4_unicode_ci —— 支持多语言排序'),
])
add_code_block(doc, '''CREATE DATABASE IF NOT EXISTS FurnitureDB
    DEFAULT CHARACTER SET utf8mb4
    DEFAULT COLLATE utf8mb4_unicode_ci;

USE FurnitureDB;''')

# ── 2.1 家具类型表 ──
doc.add_heading('2.1 家具类型表 (FurnitureCategory)', level=2)
add_explain_table(doc, [
    ('字段名', '数据类型', '约束', '说明'),
    ('CategoryID', 'INT', 'PK, AUTO_INCREMENT', '类型编号，自增主键'),
    ('CategoryName', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '类型名称（如沙发、床等）'),
    ('Description', 'VARCHAR(200)', 'DEFAULT NULL', '类型描述，可空'),
], col_widths=[4, 3, 5, 6])
add_code_block(doc, '''CREATE TABLE FurnitureCategory (
    CategoryID   INT AUTO_INCREMENT PRIMARY KEY COMMENT '家具类型编号',
    CategoryName VARCHAR(50)  NOT NULL UNIQUE COMMENT '类型名称',
    Description  VARCHAR(200) DEFAULT NULL COMMENT '类型描述'
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='家具类型表';''')

# ── 2.2 供应商表 ──
doc.add_heading('2.2 供应商信息表 (Supplier)', level=2)
add_explain_table(doc, [
    ('字段名', '数据类型', '约束', '说明'),
    ('SupplierID', 'INT', 'PK, AUTO_INCREMENT', '供应商编号'),
    ('SupplierName', 'VARCHAR(100)', 'NOT NULL', '供应商名称'),
    ('ContactPerson', 'VARCHAR(50)', 'DEFAULT NULL', '联系人姓名'),
    ('Phone', 'VARCHAR(20)', 'NOT NULL', '联系电话'),
    ('Address', 'VARCHAR(200)', 'DEFAULT NULL', '地址'),
    ('Email', 'VARCHAR(100)', 'DEFAULT NULL', '邮箱'),
], col_widths=[4, 3, 5, 6])
add_code_block(doc, '''CREATE TABLE Supplier (
    SupplierID    INT AUTO_INCREMENT PRIMARY KEY COMMENT '供应商编号',
    SupplierName  VARCHAR(100) NOT NULL COMMENT '供应商名称',
    ContactPerson VARCHAR(50)  DEFAULT NULL COMMENT '联系人',
    Phone         VARCHAR(20)  NOT NULL COMMENT '联系电话',
    Address       VARCHAR(200) DEFAULT NULL COMMENT '地址',
    Email         VARCHAR(100) DEFAULT NULL COMMENT '邮箱'
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='供应商信息表';''')

# ── 2.3 客户表 ──
doc.add_heading('2.3 客户信息表 (Customer)', level=2)
add_explain_table(doc, [
    ('字段名', '数据类型', '约束', '说明'),
    ('CustomerID', 'INT', 'PK, AUTO_INCREMENT', '客户编号'),
    ('CustomerName', 'VARCHAR(100)', 'NOT NULL', '客户名称'),
    ('Phone', 'VARCHAR(20)', 'NOT NULL', '联系电话'),
    ('Address', 'VARCHAR(200)', 'DEFAULT NULL', '地址'),
    ('Email', 'VARCHAR(100)', 'DEFAULT NULL', '邮箱'),
], col_widths=[4, 3, 5, 6])
add_code_block(doc, '''CREATE TABLE Customer (
    CustomerID   INT AUTO_INCREMENT PRIMARY KEY COMMENT '客户编号',
    CustomerName VARCHAR(100) NOT NULL COMMENT '客户名称',
    Phone        VARCHAR(20)  NOT NULL COMMENT '联系电话',
    Address      VARCHAR(200) DEFAULT NULL COMMENT '地址',
    Email        VARCHAR(100) DEFAULT NULL COMMENT '邮箱'
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='客户信息表';''')

# ── 2.4 家具信息表 ──
doc.add_heading('2.4 家具信息表 (Furniture) —— 核心表', level=2)
add_explain_table(doc, [
    ('字段名', '数据类型', '约束', '说明'),
    ('FurnitureID', 'INT', 'PK, AUTO_INCREMENT', '家具编号'),
    ('FurnitureName', 'VARCHAR(100)', 'NOT NULL', '家具名称'),
    ('CategoryID', 'INT', 'FK → FurnitureCategory', '所属类型（外键）'),
    ('SupplierID', 'INT', 'FK → Supplier', '供应商（外键）'),
    ('UnitPrice', 'DECIMAL(10,2)', 'NOT NULL', '单价'),
    ('Stock', 'INT', 'NOT NULL DEFAULT 0', '当前库存，由触发器自动维护'),
    ('Description', 'VARCHAR(500)', 'DEFAULT NULL', '家具描述'),
], col_widths=[4, 3, 5, 6])
add_code_block(doc, '''CREATE TABLE Furniture (
    FurnitureID   INT AUTO_INCREMENT PRIMARY KEY COMMENT '家具编号',
    FurnitureName VARCHAR(100) NOT NULL COMMENT '家具名称',
    CategoryID    INT          NOT NULL COMMENT '所属类型编号',
    SupplierID    INT          NOT NULL COMMENT '供应商编号',
    UnitPrice     DECIMAL(10,2) NOT NULL COMMENT '单价',
    Stock         INT          NOT NULL DEFAULT 0 COMMENT '当前库存数量',
    Description   VARCHAR(500) DEFAULT NULL COMMENT '家具描述'
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='家具信息表';''')
p = doc.add_paragraph()
run = p.add_run('★ 核心要点：Furniture 表是系统的中枢，通过两个外键（CategoryID、SupplierID）关联维度表，'
                '同时被 WarehouseEntry 和 Sales 引用。Stock 字段由触发器自动维护，无需手动更新。')
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

# ── 2.5 入库记录表 ──
doc.add_heading('2.5 入库记录表 (WarehouseEntry)', level=2)
add_explain_table(doc, [
    ('字段名', '数据类型', '约束', '说明'),
    ('EntryID', 'INT', 'PK, AUTO_INCREMENT', '入库编号'),
    ('FurnitureID', 'INT', 'FK → Furniture', '入库的家具（外键）'),
    ('EntryDate', 'DATETIME', 'NOT NULL', '入库日期'),
    ('Quantity', 'INT', 'NOT NULL', '入库数量'),
    ('UnitPrice', 'DECIMAL(10,2)', 'NOT NULL', '入库单价'),
    ('TotalAmount', 'DECIMAL(12,2)', '计算列 (STORED)', '总金额 = Quantity × UnitPrice'),
    ('Remarks', 'VARCHAR(200)', 'DEFAULT NULL', '备注'),
], col_widths=[4, 3, 5, 6])
add_code_block(doc, '''CREATE TABLE WarehouseEntry (
    EntryID     INT AUTO_INCREMENT PRIMARY KEY COMMENT '入库编号',
    FurnitureID INT          NOT NULL COMMENT '家具编号',
    EntryDate   DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '入库日期',
    Quantity    INT          NOT NULL COMMENT '入库数量',
    UnitPrice   DECIMAL(10,2) NOT NULL COMMENT '入库单价',
    TotalAmount DECIMAL(12,2) GENERATED ALWAYS AS (Quantity * UnitPrice) STORED COMMENT '入库总金额',
    Remarks     VARCHAR(200) DEFAULT NULL COMMENT '备注'
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='入库记录表';''')
p = doc.add_paragraph()
run = p.add_run('★ 技术点：TotalAmount 是 MySQL 8.0 的 GENERATED ALWAYS AS ... STORED 计算列，'
                '由数据库自动计算并持久化存储，无需在应用程序中计算。')
run.bold = True

# ── 2.6 销售记录表 ──
doc.add_heading('2.6 销售记录表 (Sales)', level=2)
add_explain_table(doc, [
    ('字段名', '数据类型', '约束', '说明'),
    ('SaleID', 'INT', 'PK, AUTO_INCREMENT', '销售编号'),
    ('FurnitureID', 'INT', 'FK → Furniture', '销售的家具（外键）'),
    ('CustomerID', 'INT', 'FK → Customer', '客户（外键）'),
    ('SaleDate', 'DATETIME', 'NOT NULL', '销售日期'),
    ('Quantity', 'INT', 'NOT NULL', '销售数量'),
    ('UnitPrice', 'DECIMAL(10,2)', 'NOT NULL', '销售单价'),
    ('TotalAmount', 'DECIMAL(12,2)', '计算列 (STORED)', '总金额 = Quantity × UnitPrice'),
    ('Remarks', 'VARCHAR(200)', 'DEFAULT NULL', '备注'),
], col_widths=[4, 3, 5, 6])
add_code_block(doc, '''CREATE TABLE Sales (
    SaleID      INT AUTO_INCREMENT PRIMARY KEY COMMENT '销售编号',
    FurnitureID INT          NOT NULL COMMENT '家具编号',
    CustomerID  INT          NOT NULL COMMENT '客户编号',
    SaleDate    DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '销售日期',
    Quantity    INT          NOT NULL COMMENT '销售数量',
    UnitPrice   DECIMAL(10,2) NOT NULL COMMENT '销售单价',
    TotalAmount DECIMAL(12,2) GENERATED ALWAYS AS (Quantity * UnitPrice) STORED COMMENT '销售总金额',
    Remarks     VARCHAR(200) DEFAULT NULL COMMENT '备注'
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='销售记录表';''')

# ── 2.7 收款记录表 ──
doc.add_heading('2.7 收款记录表 (Receipt)', level=2)
add_explain_table(doc, [
    ('字段名', '数据类型', '约束', '说明'),
    ('ReceiptID', 'INT', 'PK, AUTO_INCREMENT', '收款编号'),
    ('SaleID', 'INT', 'FK → Sales', '关联的销售记录（外键）'),
    ('ReceiptDate', 'DATETIME', 'NOT NULL', '收款日期'),
    ('Amount', 'DECIMAL(12,2)', 'NOT NULL', '收款金额'),
    ('PaymentMethod', 'VARCHAR(30)', "DEFAULT '现金'", '支付方式（现金/刷卡/转账等）'),
    ('Remarks', 'VARCHAR(200)', 'DEFAULT NULL', '备注'),
], col_widths=[4, 3, 5, 6])
add_code_block(doc, '''CREATE TABLE Receipt (
    ReceiptID     INT AUTO_INCREMENT PRIMARY KEY COMMENT '收款编号',
    SaleID        INT          NOT NULL COMMENT '关联销售编号',
    ReceiptDate   DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '收款日期',
    Amount        DECIMAL(12,2) NOT NULL COMMENT '收款金额',
    PaymentMethod VARCHAR(30)  DEFAULT '现金' COMMENT '支付方式（现金/刷卡/转账/支付宝/微信）',
    Remarks       VARCHAR(200) DEFAULT NULL COMMENT '备注'
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='收款记录表';''')

doc.add_page_break()

# ================================================================
# 三、参照完整性约束
# ================================================================
doc.add_heading('三、参照完整性约束（外键）', level=1)

doc.add_paragraph('外键约束是保证数据一致性的核心机制。本系统定义 6 个外键，'
                  '全部使用 ON DELETE RESTRICT + ON UPDATE CASCADE 策略：')

add_explain_table(doc, [
    ('编号', '外键', '引用关系', '删除策略', '更新策略'),
    ('FK1', 'Furniture.CategoryID', '→ FurnitureCategory.CategoryID', 'RESTRICT', 'CASCADE'),
    ('FK2', 'Furniture.SupplierID', '→ Supplier.SupplierID', 'RESTRICT', 'CASCADE'),
    ('FK3', 'WarehouseEntry.FurnitureID', '→ Furniture.FurnitureID', 'RESTRICT', 'CASCADE'),
    ('FK4', 'Sales.FurnitureID', '→ Furniture.FurnitureID', 'RESTRICT', 'CASCADE'),
    ('FK5', 'Sales.CustomerID', '→ Customer.CustomerID', 'RESTRICT', 'CASCADE'),
    ('FK6', 'Receipt.SaleID', '→ Sales.SaleID', 'RESTRICT', 'CASCADE'),
], col_widths=[2, 5, 5, 3, 3])

doc.add_paragraph()
doc.add_heading('策略说明', level=2)
add_explain_table(doc, [
    ('策略', '含义', '示例场景'),
    ('ON DELETE RESTRICT', '删除父记录时，如果有子记录引用，则拒绝删除', '类型下有家具时，不能删除该类型'),
    ('ON UPDATE CASCADE', '父表主键更新时，子表的外键自动同步更新', '类型ID从3改为30，所有家具的CategoryID自动变为30'),
], col_widths=[4, 5, 7])

doc.add_paragraph('完整 ALTER TABLE 语句：')
add_code_block(doc, '''-- ① 家具 → 类型
ALTER TABLE Furniture
    ADD CONSTRAINT FK_Furniture_Category
    FOREIGN KEY (CategoryID) REFERENCES FurnitureCategory(CategoryID)
    ON DELETE RESTRICT ON UPDATE CASCADE;

-- ② 家具 → 供应商
ALTER TABLE Furniture
    ADD CONSTRAINT FK_Furniture_Supplier
    FOREIGN KEY (SupplierID) REFERENCES Supplier(SupplierID)
    ON DELETE RESTRICT ON UPDATE CASCADE;

-- ③ 入库 → 家具
ALTER TABLE WarehouseEntry
    ADD CONSTRAINT FK_WarehouseEntry_Furniture
    FOREIGN KEY (FurnitureID) REFERENCES Furniture(FurnitureID)
    ON DELETE RESTRICT ON UPDATE CASCADE;

-- ④ 销售 → 家具
ALTER TABLE Sales
    ADD CONSTRAINT FK_Sales_Furniture
    FOREIGN KEY (FurnitureID) REFERENCES Furniture(FurnitureID)
    ON DELETE RESTRICT ON UPDATE CASCADE;

-- ⑤ 销售 → 客户
ALTER TABLE Sales
    ADD CONSTRAINT FK_Sales_Customer
    FOREIGN KEY (CustomerID) REFERENCES Customer(CustomerID)
    ON DELETE RESTRICT ON UPDATE CASCADE;

-- ⑥ 收款 → 销售
ALTER TABLE Receipt
    ADD CONSTRAINT FK_Receipt_Sales
    FOREIGN KEY (SaleID) REFERENCES Sales(SaleID)
    ON DELETE RESTRICT ON UPDATE CASCADE;''')

doc.add_page_break()

# ================================================================
# 四、触发器
# ================================================================
doc.add_heading('四、触发器（自动更新库存）', level=1)

doc.add_paragraph('触发器在 INSERT / DELETE / UPDATE 操作后自动执行，本系统通过 6 个触发器实现库存的自动维护。'
                  '全部使用 AFTER 时机，确保数据已写入成功后才修改库存。')

doc.add_heading('触发器逻辑总览', level=2)
add_explain_table(doc, [
    ('表', '操作', '触发器名', 'SQL逻辑', '说明'),
    ('WarehouseEntry', 'INSERT', 'trg_WarehouseEntry_AfterInsert', 'Stock += NEW.Quantity', '入库 → 库存增加'),
    ('WarehouseEntry', 'DELETE', 'trg_WarehouseEntry_AfterDelete', 'Stock -= OLD.Quantity', '删入库 → 库存回退'),
    ('WarehouseEntry', 'UPDATE', 'trg_WarehouseEntry_AfterUpdate', 'Stock = Stock - OLD.Qty + NEW.Qty', '改入库 → 调整库存'),
    ('Sales', 'INSERT', 'trg_Sales_AfterInsert', 'Stock -= NEW.Quantity', '销售 → 库存减少'),
    ('Sales', 'DELETE', 'trg_Sales_AfterDelete', 'Stock += OLD.Quantity', '删销售 → 库存还原'),
    ('Sales', 'UPDATE', 'trg_Sales_AfterUpdate', 'Stock = Stock + OLD.Qty - NEW.Qty', '改销售 → 调整库存'),
], col_widths=[2.5, 2, 4, 4, 3.5])

doc.add_heading('NEW / OLD 伪行机制', level=2)
add_explain_table(doc, [
    ('操作类型', '可用伪行', '说明'),
    ('INSERT', '只有 NEW', 'NEW 包含即将插入的新数据'),
    ('DELETE', '只有 OLD', 'OLD 包含已被删除的旧数据'),
    ('UPDATE', '同时有 NEW 和 OLD', 'NEW 是更新后的值，OLD 是更新前的值'),
], col_widths=[3, 3, 9])

doc.add_paragraph()
doc.add_heading('4.1 入库触发器完整代码', level=2)
add_code_block(doc, '''DELIMITER //

-- 入库 INSERT 触发器
CREATE TRIGGER trg_WarehouseEntry_AfterInsert
AFTER INSERT ON WarehouseEntry
FOR EACH ROW
BEGIN
    UPDATE Furniture
    SET Stock = Stock + NEW.Quantity
    WHERE FurnitureID = NEW.FurnitureID;
END//

-- 入库 DELETE 触发器
CREATE TRIGGER trg_WarehouseEntry_AfterDelete
AFTER DELETE ON WarehouseEntry
FOR EACH ROW
BEGIN
    UPDATE Furniture
    SET Stock = Stock - OLD.Quantity
    WHERE FurnitureID = OLD.FurnitureID;
END//

-- 入库 UPDATE 触发器
CREATE TRIGGER trg_WarehouseEntry_AfterUpdate
AFTER UPDATE ON WarehouseEntry
FOR EACH ROW
BEGIN
    UPDATE Furniture
    SET Stock = Stock - OLD.Quantity + NEW.Quantity
    WHERE FurnitureID = NEW.FurnitureID;
END//

DELIMITER ;''')

doc.add_paragraph()
doc.add_heading('4.2 销售触发器完整代码', level=2)
add_code_block(doc, '''DELIMITER //

-- 销售 INSERT 触发器
CREATE TRIGGER trg_Sales_AfterInsert
AFTER INSERT ON Sales
FOR EACH ROW
BEGIN
    UPDATE Furniture
    SET Stock = Stock - NEW.Quantity
    WHERE FurnitureID = NEW.FurnitureID;
END//

-- 销售 DELETE 触发器
CREATE TRIGGER trg_Sales_AfterDelete
AFTER DELETE ON Sales
FOR EACH ROW
BEGIN
    UPDATE Furniture
    SET Stock = Stock + OLD.Quantity
    WHERE FurnitureID = OLD.FurnitureID;
END//

-- 销售 UPDATE 触发器
CREATE TRIGGER trg_Sales_AfterUpdate
AFTER UPDATE ON Sales
FOR EACH ROW
BEGIN
    UPDATE Furniture
    SET Stock = Stock + OLD.Quantity - NEW.Quantity
    WHERE FurnitureID = NEW.FurnitureID;
END//

DELIMITER ;''')

doc.add_paragraph()
doc.add_heading('扩展：库存不足保护（BEFORE 触发器）', level=2)
doc.add_paragraph('如果需要数据库层面阻止超卖，可以添加一个 BEFORE INSERT 触发器，在销售写入前检查库存是否充足：')
add_code_block(doc, '''DELIMITER //

CREATE TRIGGER trg_Sales_BeforeInsert
BEFORE INSERT ON Sales
FOR EACH ROW
BEGIN
    DECLARE cur_stock INT;
    
    -- 获取当前库存
    SELECT Stock INTO cur_stock
    FROM Furniture
    WHERE FurnitureID = NEW.FurnitureID;
    
    -- 库存不足则报错，阻止写入
    IF cur_stock < NEW.Quantity THEN
        SIGNAL SQLSTATE '45000'
        SET MESSAGE_TEXT = '库存不足，无法完成销售';
    END IF;
END//

DELIMITER ;''')

doc.add_page_break()

# ================================================================
# 五、存储过程
# ================================================================
doc.add_heading('五、存储过程（时间段统计）', level=1)

doc.add_paragraph('存储过程 sp_InventoryStatistics 接收开始和结束时间两个参数，'
                  '统计指定时间段内每种家具的入库总数量和销售总数量。')

doc.add_heading('参数说明', level=2)
add_explain_table(doc, [
    ('参数名', '数据类型', '方向', '说明'),
    ('startDate', 'DATETIME', 'IN', '统计开始时间'),
    ('endDate', 'DATETIME', 'IN', '统计结束时间'),
], col_widths=[3, 3, 2, 7])

doc.add_paragraph()
doc.add_heading('SQL 技术要点', level=2)
add_explain_table(doc, [
    ('技术点', '说明', '代码片段'),
    ('子查询(派生表)', '先按 FurnitureID 汇总入库和销售数量，再与主表 JOIN',
     '(SELECT FurnitureID, SUM(Quantity) AS in_qty FROM WarehouseEntry WHERE ... GROUP BY FurnitureID) in_stats'),
    ('GROUP BY 聚合', '按 FurnitureID 分组，SUM 求和', 'GROUP BY FurnitureID'),
    ('LEFT JOIN + IFNULL', '没有入库/销售的家具也显示，显示为 0',
     'LEFT JOIN ... ON ... IFNULL(in_qty, 0) AS total_in'),
    ('BETWEEN 筛选', '筛选指定时间范围内的记录', 'WHERE EntryDate BETWEEN startDate AND endDate'),
], col_widths=[3, 5, 7])

doc.add_paragraph()
doc.add_heading('存储过程完整代码', level=2)
add_code_block(doc, '''DELIMITER //

CREATE PROCEDURE sp_InventoryStatistics(
    IN startDate DATETIME,
    IN endDate   DATETIME
)
BEGIN
    SELECT
        f.FurnitureID          AS '家具编号',
        f.FurnitureName        AS '家具名称',
        fc.CategoryName        AS '家具类型',
        s.SupplierName         AS '供应商',
        IFNULL(in_stats.TotalInQty, 0)    AS '入库总数量',
        IFNULL(sale_stats.TotalSaleQty, 0) AS '销售总数量',
        f.UnitPrice            AS '当前单价',
        f.Stock                AS '当前库存'
    FROM Furniture f
    LEFT JOIN FurnitureCategory fc ON f.CategoryID = fc.CategoryID
    LEFT JOIN Supplier s ON f.SupplierID = s.SupplierID
    LEFT JOIN (
        SELECT FurnitureID, SUM(Quantity) AS TotalInQty
        FROM WarehouseEntry
        WHERE EntryDate BETWEEN startDate AND endDate
        GROUP BY FurnitureID
    ) in_stats ON f.FurnitureID = in_stats.FurnitureID
    LEFT JOIN (
        SELECT FurnitureID, SUM(Quantity) AS TotalSaleQty
        FROM Sales
        WHERE SaleDate BETWEEN startDate AND endDate
        GROUP BY FurnitureID
    ) sale_stats ON f.FurnitureID = sale_stats.FurnitureID
    ORDER BY f.FurnitureID;
END//

DELIMITER ;''')

doc.add_paragraph()
doc.add_heading('调用示例', level=2)
add_code_block(doc, '-- 查询 2024年1月1日 到 2024年3月31日 的统计\n'
                    'CALL sp_InventoryStatistics(\'2024-01-01 00:00:00\', \'2024-03-31 23:59:59\');\n\n'
                    '-- 返回结果示例:\n'
                    '+----------+------------------+----------+------------+----------+----------+----------+----------+\n'
                    '| 家具编号 | 家具名称         | 家具类型 | 供应商     | 入库总数 | 销售总数 | 当前单价 | 当前库存 |\n'
                    '+----------+------------------+----------+------------+----------+----------+----------+----------+\n'
                    '|    1     | 真皮三人沙发      | 沙发     | 光明家具厂 |   20     |   1      | 3800.00  |   19    |\n'
                    '|    2     | 布艺双人沙发      | 沙发     | 红星家具   |   25     |   5      | 2200.00  |   20    |\n'
                    '+----------+------------------+----------+------------+----------+----------+----------+----------+')

doc.add_page_break()

# ================================================================
# 六、视图
# ================================================================
doc.add_heading('六、视图封装（简化查询）', level=1)

doc.add_paragraph('视图将复杂的 3~4 表 JOIN 封装为虚拟表，应用程序查询时像使用单表一样简单。'
                  '本系统定义了 3 个常用视图：')

doc.add_heading('6.1 家具库存视图 (v_FurnitureStock)', level=2)
add_explain_table(doc, [
    ('视图名', 'v_FurnitureStock'),
    ('关联表数', '3 (Furniture + FurnitureCategory + Supplier)'),
    ('用途', '查看所有家具的完整信息（含类型名称、供应商名称、库存数量）'),
])
add_code_block(doc, '''CREATE OR REPLACE VIEW v_FurnitureStock AS
SELECT
    f.FurnitureID,
    f.FurnitureName,
    fc.CategoryName,
    s.SupplierName,
    f.UnitPrice,
    f.Stock,
    f.Description
FROM Furniture f
JOIN FurnitureCategory fc ON f.CategoryID = fc.CategoryID
JOIN Supplier s ON f.SupplierID = s.SupplierID;

-- 使用示例：查看库存不足5件的家具
SELECT * FROM v_FurnitureStock WHERE Stock < 5 ORDER BY Stock;''')

doc.add_paragraph()
doc.add_heading('6.2 销售明细视图 (v_SalesDetail)', level=2)
add_explain_table(doc, [
    ('视图名', 'v_SalesDetail'),
    ('关联表数', '3 (Sales + Furniture + Customer)'),
    ('用途', '查看销售记录的完整信息，含家具名称和客户名称'),
])
add_code_block(doc, '''CREATE OR REPLACE VIEW v_SalesDetail AS
SELECT
    sa.SaleID,
    f.FurnitureName,
    c.CustomerName,
    sa.SaleDate,
    sa.Quantity,
    sa.UnitPrice,
    sa.TotalAmount,
    sa.Remarks
FROM Sales sa
JOIN Furniture f ON sa.FurnitureID = f.FurnitureID
JOIN Customer c ON sa.CustomerID = c.CustomerID;

-- 使用示例：查看某个客户的购买记录
SELECT * FROM v_SalesDetail WHERE CustomerName = '李明';''')

doc.add_paragraph()
doc.add_heading('6.3 收款明细视图 (v_ReceiptDetail)', level=2)
add_explain_table(doc, [
    ('视图名', 'v_ReceiptDetail'),
    ('关联表数', '4 (Receipt + Sales + Furniture + Customer)'),
    ('用途', '查看收款记录的完整信息，含家具名称、客户名称、支付方式'),
])
add_code_block(doc, '''CREATE OR REPLACE VIEW v_ReceiptDetail AS
SELECT
    r.ReceiptID,
    f.FurnitureName,
    c.CustomerName,
    sa.SaleDate,
    r.ReceiptDate,
    r.Amount,
    r.PaymentMethod,
    r.Remarks
FROM Receipt r
JOIN Sales sa ON r.SaleID = sa.SaleID
JOIN Furniture f ON sa.FurnitureID = f.FurnitureID
JOIN Customer c ON sa.CustomerID = c.CustomerID;

-- 使用示例
SELECT * FROM v_ReceiptDetail WHERE PaymentMethod = '转账';''')

doc.add_page_break()

# ================================================================
# 七、测试数据
# ================================================================
doc.add_heading('七、测试数据（预置记录）', level=1)

doc.add_paragraph('脚本运行后自动插入以下测试数据，可用于验证所有功能：')

add_explain_table(doc, [
    ('表名', '记录数', '数据说明'),
    ('FurnitureCategory', '6', '沙发、床、餐桌、书桌、柜子、椅子'),
    ('Supplier', '4', '光明家具厂、红星家具集团、大自然家居、北欧风尚家具'),
    ('Customer', '5', '个人客户3个 + 企业客户2个（四季酒店、阳光地产）'),
    ('Furniture', '10', '含真皮沙发、布艺沙发、实木床、餐桌、书桌、衣柜、办公椅等'),
    ('WarehouseEntry', '10', '2024年1~2月的入库记录，每件家具至少入库一次'),
    ('Sales', '10', '2024年1~3月的销售记录，覆盖多位客户'),
    ('Receipt', '10', '每笔销售对应一条收款记录，含现金、刷卡、转账、微信、支付宝'),
], col_widths=[4, 2, 10])

doc.add_page_break()

# ================================================================
# 八、Web 应用功能说明（Flask 前端特性）
# ================================================================
doc.add_heading('八、Web 应用功能说明（Flask 前端特性）', level=1)

doc.add_paragraph('本系统配套的 Flask Web 应用提供了完整的图形化操作界面，核心特性如下：')

doc.add_heading('8.1 弹窗内联创建 — 减少操作步骤', level=2)
doc.add_paragraph('在入库和销售操作时，支持弹窗内联创建关联数据，无需提前创建或多次弹窗：')

add_explain_table(doc, [
    ('操作页面', '可内联创建的对象', '操作方式', '前后对比'),
    ('入库管理', '供应商 + 家具', '弹窗内展开面板，填写后 AJAX 保存并自动选中',
     '3次弹窗 → 1次弹窗完成'),
    ('销售管理', '客户 + 供应商 + 家具', '弹窗内展开面板，填写后 AJAX 保存并自动选中',
     '4次弹窗 → 1次弹窗完成'),
], col_widths=[3, 3.5, 5, 3])

doc.add_paragraph()
doc.add_heading('8.2 库存自动更新机制', level=2)
doc.add_paragraph('由于 SQLite 不支持 MySQL 的跨表触发器，Web 应用在 Python 代码层实现了等价的库存自动更新逻辑：')

add_explain_table(doc, [
    ('操作', 'Python 代码逻辑', '等价于 MySQL 触发器'),
    ('新增入库', 'INSERT WarehouseEntry → UPDATE Furniture SET Stock = Stock + Qty', 'trg_WarehouseEntry_AfterInsert'),
    ('删除入库', 'DELETE WarehouseEntry → UPDATE Furniture SET Stock = Stock - Qty', 'trg_WarehouseEntry_AfterDelete'),
    ('新增销售', '先检查 Stock >= Qty → INSERT Sales → UPDATE Stock = Stock - Qty', 'trg_Sales_AfterInsert + 校验'),
    ('删除销售', 'DELETE Sales → UPDATE Furniture SET Stock = Stock + Qty', 'trg_Sales_AfterDelete'),
], col_widths=[3, 6, 5])

doc.add_paragraph()
doc.add_heading('8.3 页面路由一览', level=2)
add_explain_table(doc, [
    ('URL路径', '功能', '技术要点'),
    ('/', '仪表盘', '系统数据总览 + 低库存预警'),
    ('/categories', '家具类型管理', '增删列表，外键保护'),
    ('/suppliers', '供应商管理', '增删，联系方式'),
    ('/customers', '客户管理', '增删，联系方式'),
    ('/furniture', '家具信息管理', '关联类型+供应商，库存显示'),
    ('/warehouse', '入库管理', '内联创建 + 库存自动更新'),
    ('/sales', '销售管理', '库存校验 + 内联创建 + 自动扣减'),
    ('/receipts', '收款管理', '关联销售 + AJAX实时查询应收/已收/未收 + 超收预警'),
    ('/stock', '库存盘点', '三色状态标（充足/偏低/告急）'),
    ('/statistics', '统计分析', '时间段筛选，入库/销售对比'),
], col_widths=[3, 4, 7])

doc.add_page_break()

# ================================================================
# 九、常用查询示例
# ================================================================
doc.add_heading('九、常用查询示例', level=1)

doc.add_paragraph('以下汇总了本系统最常用的查询语句，覆盖了 3 表 JOIN、4 表 JOIN、聚合统计等典型场景：')

queries = [
    ('查询所有家具及其类型、供应商信息',
     'SELECT f.FurnitureID, f.FurnitureName, c.CategoryName, s.SupplierName, f.UnitPrice, f.Stock\n'
     'FROM Furniture f\n'
     'JOIN FurnitureCategory c ON f.CategoryID = c.CategoryID\n'
     'JOIN Supplier s ON f.SupplierID = s.SupplierID\n'
     'ORDER BY f.FurnitureID;'),
    ('查询低库存商品（库存 < 5）',
     'SELECT * FROM v_FurnitureStock WHERE Stock < 5 ORDER BY Stock;'),
    ('查询某客户的购买记录',
     'SELECT * FROM v_SalesDetail WHERE CustomerName = \'李明\';'),
    ('查询某供应商供应的所有家具',
     'SELECT f.FurnitureName, f.UnitPrice, f.Stock\n'
     'FROM Furniture f\n'
     'JOIN Supplier s ON f.SupplierID = s.SupplierID\n'
     'WHERE s.SupplierName = \'光明家具厂\';'),
    ('统计各类型家具的数量和平均价格',
     'SELECT c.CategoryName, COUNT(*) AS 家具数量, ROUND(AVG(f.UnitPrice), 2) AS 平均价格\n'
     'FROM Furniture f\n'
     'JOIN FurnitureCategory c ON f.CategoryID = c.CategoryID\n'
     'GROUP BY c.CategoryName\n'
     'ORDER BY 家具数量 DESC;'),
    ('查询某段时间内的入库笔数和总量',
     'SELECT COUNT(*) AS 入库笔数, SUM(Quantity) AS 入库总量\n'
     'FROM WarehouseEntry\n'
     'WHERE EntryDate BETWEEN \'2024-01-01\' AND \'2024-01-31\';'),
    ('查询每种家具的总销售额（聚合）',
     'SELECT f.FurnitureName, SUM(s.Quantity * s.UnitPrice) AS 总销售额\n'
     'FROM Sales s\n'
     'JOIN Furniture f ON s.FurnitureID = f.FurnitureID\n'
     'GROUP BY f.FurnitureName\n'
     'ORDER BY 总销售额 DESC;'),
    ('查询所有未收全款的销售（应收 vs 已收）',
     'SELECT s.SaleID, f.FurnitureName, c.CustomerName,\n'
     '       s.TotalAmount AS 应收金额,\n'
     '       COALESCE(SUM(r.Amount), 0) AS 已收金额,\n'
     '       s.TotalAmount - COALESCE(SUM(r.Amount), 0) AS 未收金额\n'
     'FROM Sales s\n'
     'JOIN Furniture f ON s.FurnitureID = f.FurnitureID\n'
     'JOIN Customer c ON s.CustomerID = c.CustomerID\n'
     'LEFT JOIN Receipt r ON s.SaleID = r.SaleID\n'
     'GROUP BY s.SaleID, f.FurnitureName, c.CustomerName, s.TotalAmount\n'
     'HAVING 未收金额 > 0;'),
]

for i, (desc, sql) in enumerate(queries, 1):
    doc.add_heading(f'查询{i}：{desc}', level=2)
    add_code_block(doc, sql, font_size=10)

# ================================================================
# 保存
# ================================================================
output_path = os.path.join(os.path.dirname(__file__), '家具城进销存管理系统_SQL实现文档.docx')
# If old file is locked, fallback
try:
    doc.save(output_path)
    print(f'文档已保存: {output_path}')
except PermissionError:
    alt = os.path.join(os.path.dirname(__file__), 'FurnitureDB_SQL_Doc.docx')
    doc.save(alt)
    print(f'文档已保存(备用名): {alt}')
